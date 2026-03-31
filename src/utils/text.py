from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List


KB_SUPPORTED_EXTENSIONS = (
	".txt",
	".md",
	".markdown",
	".json",
	".csv",
	".docx",
	".pdf",
)


def discover_text_files(root: str | os.PathLike[str]) -> List[Path]:
	"""Recursively discover KB source files supported by ingest."""
	root_path = Path(root)
	candidates: List[Path] = []
	for path in root_path.rglob("*"):
		if path.is_file() and path.suffix.lower() in KB_SUPPORTED_EXTENSIONS:
			candidates.append(path)
	return sorted(candidates)


def _read_docx_text(path: Path) -> str:
	try:
		from docx import Document  # type: ignore
	except Exception as exc:
		raise RuntimeError("读取 .docx 需要安装 python-docx（pip install python-docx）") from exc

	doc = Document(str(path))
	lines: list[str] = []
	for p in doc.paragraphs:
		text = (p.text or "").strip()
		if text:
			lines.append(text)
	for table in doc.tables:
		for row in table.rows:
			cells = [(c.text or "").strip() for c in row.cells]
			cells = [c for c in cells if c]
			if cells:
				lines.append(" | ".join(cells))
	return "\n".join(lines)


def _read_pdf_text(path: Path) -> str:
	reader_cls = None
	try:
		from pypdf import PdfReader as reader_cls  # type: ignore
	except Exception:
		try:
			from PyPDF2 import PdfReader as reader_cls  # type: ignore
		except Exception as exc:
			raise RuntimeError("读取 .pdf 需要安装 pypdf 或 PyPDF2（pip install pypdf）") from exc

	reader = reader_cls(str(path))
	parts: list[str] = []
	for page in reader.pages:
		text = (page.extract_text() or "").strip()
		if text:
			parts.append(text)
	return "\n\n".join(parts)


def read_file_text(path: Path) -> str:
	ext = path.suffix.lower()
	if ext == ".docx":
		return _read_docx_text(path)
	if ext == ".pdf":
		return _read_pdf_text(path)
	with path.open("r", encoding="utf-8", errors="ignore") as f:
		return f.read()


def clean_text(text: str) -> str:
	# Normalize line breaks and collapse excessive whitespace.
	text = text.replace("\r\n", "\n").replace("\r", "\n")
	text = re.sub("\u3000", " ", text)  # full-width space
	text = re.sub("\n{3,}", "\n\n", text)
	return text.strip()


def split_by_length(text: str, max_chars: int = 800, overlap: int = 120) -> List[str]:
	"""Character-based splitter that tries to respect sentence boundaries."""
	if not text:
		return []
	sentences = re.split(r"(?<=[。！？!?.])\s+", text)
	chunks: List[str] = []
	current: List[str] = []
	current_len = 0
	for s in sentences:
		s_len = len(s)
		if current_len + s_len <= max_chars or not current:
			current.append(s)
			current_len += s_len
		else:
			chunks.append("".join(current).strip())
			# Add overlap by taking tail of previous chunk
			if overlap > 0 and chunks[-1]:
				tail = chunks[-1][-overlap:]
				current = [tail, s]
				current_len = len(tail) + s_len
			else:
				current = [s]
				current_len = s_len
	if current:
		chunks.append("".join(current).strip())
	# Filter tiny chunks
	return [c for c in chunks if len(c) > 20]


def sanitize(s: str) -> str:
	"""清理字符串中的特殊字符和前缀"""
	# Remove quotes, bearer prefix, and exotic spaces
	val = (s or "")
	val = val.replace("\u200b", "").replace("\u2003", " ").replace("\u00A0", " ")
	val = val.strip().strip('"').strip("'")
	if val.lower().startswith("bearer "):
		val = val[7:].strip()
	return val


def _stringify_chat_content(value) -> str:
	"""Convert chat content payloads to plain text."""
	if value is None:
		return ""
	if isinstance(value, str):
		return value.strip()
	if isinstance(value, (list, tuple)):
		parts: list[str] = []
		for item in value:
			if isinstance(item, str):
				parts.append(item)
				continue
			if isinstance(item, dict):
				parts.append(str(item.get("text") or item.get("content") or "").strip())
				continue
			text = getattr(item, "text", None) or getattr(item, "content", None)
			if text:
				parts.append(str(text).strip())
		return " ".join([p for p in parts if p]).strip()
	if isinstance(value, dict):
		return str(value.get("text") or value.get("content") or value.get("message") or "").strip()
	return str(value).strip()


def _extract_chat_reply(resp) -> str:
	"""Extract a readable reply from heterogeneous API response payloads."""
	if resp is None:
		return ""
	if isinstance(resp, str):
		return resp.strip()

	if isinstance(resp, dict):
		choices = resp.get("choices")
		if isinstance(choices, list) and choices:
			first = choices[0]
			if isinstance(first, dict):
				message = first.get("message")
				if isinstance(message, dict):
					reply = _stringify_chat_content(message.get("content") or message.get("text"))
				else:
					reply = _stringify_chat_content(first.get("content") or first.get("text") or message)
				if reply:
					return reply
		for key in ("message", "content", "text", "output_text", "response"):
			reply = _stringify_chat_content(resp.get(key))
			if reply:
				return reply

	choices = getattr(resp, "choices", None)
	if isinstance(choices, list) and choices:
		first = choices[0]
		if isinstance(first, dict):
			message = first.get("message")
			if isinstance(message, dict):
				reply = _stringify_chat_content(message.get("content") or message.get("text"))
			else:
				reply = _stringify_chat_content(first.get("content") or first.get("text") or message)
		else:
			message = getattr(first, "message", None)
			reply = _stringify_chat_content(getattr(message, "content", None) if message is not None else None)
			if not reply:
				reply = _stringify_chat_content(getattr(message, "text", None) if message is not None else None)
			if not reply:
				reply = _stringify_chat_content(getattr(first, "text", None) or getattr(first, "content", None))
		if reply:
			return reply

	for attr in ("output_text", "text", "content", "message"):
		reply = _stringify_chat_content(getattr(resp, attr, None))
		if reply:
			return reply

	return _stringify_chat_content(resp)


def _normalize_openai_base_url(base_url: str) -> str:
	"""Normalize OpenAI-compatible base URL to API root."""
	base = (base_url or "").strip().rstrip("/")
	if not base:
		return ""
	if base.endswith("/chat/completions"):
		base = base[: -len("/chat/completions")]
	return base.rstrip("/")


def try_chat_api(key: str, base_url: str, model: str) -> tuple[bool, str]:
	"""测试聊天API是否可用。"""
	try:
		from src.clients.custom_openai_client import create_compatible_client

		normalized_base_url = _normalize_openai_base_url(base_url)
		client = create_compatible_client(api_key=key, base_url=normalized_base_url, timeout=30)
		resp = client.chat.completions.create(
			model=model,
			messages=[{"role": "user", "content": "ping"}],
			max_tokens=5
		)
		reply = _extract_chat_reply(resp)
		if reply:
			return True, f"连接成功! 模型回复: {reply}"
		return True, "连接成功! 模型已返回响应。"
	except Exception as e:
		error_msg = str(e)

		if "blocked" in error_msg.lower():
			return False, f"请求被阻止: {error_msg}\n💡 可能原因: API Key无效、IP限制、或模型名称错误"
		elif "not found" in error_msg.lower() or "404" in error_msg:
			return False, f"未找到资源: {error_msg}\n💡 可能原因: Base URL路径错误或模型不存在"
		elif "unauthorized" in error_msg.lower() or "401" in error_msg:
			return False, f"认证失败: {error_msg}\n💡 可能原因: API Key格式错误或未激活"
		elif "timeout" in error_msg.lower():
			return False, f"连接超时: {error_msg}\n💡 可能原因: 网络问题或服务器响应慢"
		else:
			return False, error_msg

def try_image_api(key: str, base_url: str, model: str) -> tuple[bool, str]:
	"""测试图片生成API是否可用"""
	try:
		from src.clients.custom_openai_client import create_compatible_client
		
		normalized_base_url = _normalize_openai_base_url(base_url)
		client = create_compatible_client(api_key=key, base_url=normalized_base_url, timeout=30)
		
		# 根据模型选择合适的测试参数
		if "dall-e-2" in model.lower():
			# DALL-E-2 建议用较小的尺寸测试，速度快且省钱
			size = "512x512"
		else:
			# DALL-E-3 和其他模型用标准尺寸
			size = "1024x1024"
		
		# 尝试生成一个简单的测试图片（不强制要求格式）
		resp = client.images.generate(
			model=model,
			prompt="a simple test image",
			n=1,
			size=size
		)
		_ok = bool(resp.data and len(resp.data) > 0)
		
		# 检查返回的数据类型
		has_url = hasattr(resp.data[0], 'url') and resp.data[0].url
		has_b64 = hasattr(resp.data[0], 'b64_json') and resp.data[0].b64_json
		
		if has_url:
			return True, f"图片API测试成功 (返回URL格式，尺寸:{size})"
		elif has_b64:
			return True, f"图片API测试成功 (返回Base64格式，尺寸:{size})"
		else:
			return True, f"图片API测试成功 (尺寸:{size})"
	except Exception as e:
		return False, str(e)


def estimate_chars_from_outline(outline: str) -> int:
	"""根据大纲估算总字数"""
	lines = [ln.strip() for ln in outline.splitlines() if ln.strip()]
	sections = []
	for line in lines:
		# 匹配：1. XXX (1000字) 或 一、XXX (1000字)
		match = re.search(r'\((\d+)字\)', line)
		if match:
			sections.append(int(match.group(1)))
	
	if not sections:
		# 如果没有标注字数，按默认每段1000字估算
		return len([ln for ln in lines if re.match(r'^[\d一二三四五六七八九十]+[、\.]', ln)]) * 1000
	
	return sum(sections)


def parse_outline_sections(outline: str) -> list[dict[str, str]]:
	"""解析大纲为章节列表"""
	lines = [ln.strip() for ln in outline.splitlines() if ln.strip()]
	sections = []
	
	for line in lines:
		# 匹配序号开头的行
		match = re.match(r'^([\d一二三四五六七八九十]+)[、\.](.+)', line)
		if match:
			num, title = match.groups()
			
			# 提取字数（如果有）
			chars = 1000  # 默认字数
			chars_match = re.search(r'\((\d+)字\)', title)
			if chars_match:
				chars = int(chars_match.group(1))
				# 去掉字数标注
				title = re.sub(r'\s*\(\d+字\)', '', title).strip()
			
			sections.append({
				'num': num,
				'title': title.strip(),
				'chars': chars
			})
	
	return sections
